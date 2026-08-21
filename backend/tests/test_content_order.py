from __future__ import annotations

from pathlib import Path

from docx import Document

from backend.app.analyzer.chunking import build_document_chunks
from backend.app.analyzer.models import TenderChunkAnalysis
from backend.app.analyzer.tender_analyzer import TenderAnalyzer
from backend.app.llm.fake import FakeLLMProvider
from backend.app.models import ContentBlockReference, EvidenceReference, Paragraph, ParsedDocument, RequirementItem, Table
from backend.app.parser.docx_parser import parse_docx
from backend.app.parser.legacy_doc_parser import parse_legacy_doc


def test_parsed_document_defaults_to_empty_content_order() -> None:
    assert ParsedDocument(filename="synthetic", file_type="docx").content_order == []


def test_docx_preserves_paragraph_table_paragraph_body_order(tmp_path: Path) -> None:
    path = tmp_path / "interleaved.docx"
    source = Document()
    source.add_paragraph("first")
    source.add_table(rows=1, cols=1).cell(0, 0).text = "table"
    source.add_paragraph("last")
    source.save(path)

    parsed = parse_docx(path, path.name)
    assert [(item.type, item.index) for item in parsed.content_order] == [
        ("paragraph", 0),
        ("table", 0),
        ("paragraph", 1),
    ]
    assert parsed.content_order[0].index == parsed.paragraphs[0].index
    assert parsed.content_order[1].index == parsed.tables[0].index


def test_docx_preserves_multiple_interleaved_blocks(tmp_path: Path) -> None:
    path = tmp_path / "multiple.docx"
    source = Document()
    source.add_paragraph("p0")
    source.add_table(rows=1, cols=1).cell(0, 0).text = "t0"
    source.add_paragraph("p1")
    source.add_table(rows=1, cols=1).cell(0, 0).text = "t1"
    source.save(path)

    parsed = parse_docx(path, path.name)
    assert [(item.type, item.index) for item in parsed.content_order] == [
        ("paragraph", 0), ("table", 0), ("paragraph", 1), ("table", 1)
    ]


def test_legacy_doc_inherits_docx_content_order(tmp_path: Path) -> None:
    source = tmp_path / "legacy.doc"
    source.write_bytes(b"synthetic")

    class Converter:
        def convert(self, _source: Path, output_dir: Path, _profile_dir: Path) -> Path:
            converted = output_dir / "legacy.docx"
            docx = Document()
            docx.add_paragraph("p0")
            docx.add_table(rows=1, cols=1).cell(0, 0).text = "t0"
            docx.add_paragraph("p1")
            docx.save(converted)
            return converted

    parsed = parse_legacy_doc(source, source.name, Converter())
    assert [(item.type, item.index) for item in parsed.content_order] == [
        ("paragraph", 0), ("table", 0), ("paragraph", 1)
    ]


def test_chunker_uses_content_order_instead_of_tables_last() -> None:
    parsed = ParsedDocument(
        filename="synthetic",
        file_type="docx",
        paragraphs=[Paragraph(text="p0", index=0), Paragraph(text="p1", index=1)],
        tables=[Table(index=0, rows=[["t0"]])],
        content_order=[
            ContentBlockReference(type="paragraph", index=0),
            ContentBlockReference(type="table", index=0),
            ContentBlockReference(type="paragraph", index=1),
        ],
    )
    content = build_document_chunks(parsed, max_chars=200, overlap_chars=0)[0].content
    assert content.index("[P:0") < content.index("[T:0") < content.index("[P:1")


def test_empty_content_order_falls_back_and_signals_warning() -> None:
    parsed = ParsedDocument(
        filename="synthetic",
        file_type="pdf",
        paragraphs=[Paragraph(text="p0", index=0)],
        tables=[Table(index=0, rows=[["t0"]])],
    )
    content = build_document_chunks(parsed, max_chars=200, overlap_chars=0)[0].content
    assert content.index("[P:0") < content.index("[T:0")
    result = TenderAnalyzer(FakeLLMProvider(TenderChunkAnalysis())).analyze(parsed)
    assert "source_order_unavailable" in result.warnings


def test_content_order_keeps_locator_and_evidence_validation_after_split() -> None:
    parsed = ParsedDocument(
        filename="synthetic",
        file_type="docx",
        paragraphs=[Paragraph(text="must satisfy " + "x" * 60, page_number=2, index=7)],
        content_order=[ContentBlockReference(type="paragraph", index=7)],
    )
    chunks = build_document_chunks(parsed, max_chars=30, overlap_chars=0)
    assert all(chunk.paragraph_indices == [7] for chunk in chunks)
    response = TenderChunkAnalysis(
        qualification_requirements=[
            RequirementItem(text="must", evidence=[EvidenceReference(paragraph_index=7, page_number=2, quote="must satisfy")])
        ]
    )
    result = TenderAnalyzer(FakeLLMProvider(response)).analyze(parsed)
    assert result.qualification_requirements[0].evidence[0].paragraph_index == 7
