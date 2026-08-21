from __future__ import annotations

import shutil
from pathlib import Path

import pymupdf
import pytest
from docx import Document
from fastapi.testclient import TestClient

from backend.app import main
from backend.app.analyzer.models import TenderChunkAnalysis
from backend.app.analyzer.tender_analyzer import TenderAnalyzerError
from backend.app.llm.exceptions import LLMConfigurationError, LLMConnectionError, LLMHTTPError, LLMStructuredOutputError, LLMTimeoutError
from backend.app.llm.fake import FakeLLMProvider
from backend.app.models import EvidenceReference, Paragraph, ParsedDocument, RequirementItem, TenderAnalysis
from backend.app.parser.pdf_parser import DocumentParseError


@pytest.fixture(autouse=True)
def clear_dependency_overrides() -> None:
    main.app.dependency_overrides.clear()
    yield
    main.app.dependency_overrides.clear()


@pytest.fixture()
def sample_docx(tmp_path: Path) -> Path:
    path = tmp_path / "tender.docx"
    document = Document()
    document.add_paragraph("The bidder must qualify.")
    document.save(path)
    return path


@pytest.fixture()
def sample_pdf(tmp_path: Path) -> Path:
    path = tmp_path / "tender.pdf"
    pdf = pymupdf.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "The bidder must qualify.")
    pdf.save(path)
    pdf.close()
    return path


def analysis_response() -> TenderChunkAnalysis:
    return TenderChunkAnalysis(
        qualification_requirements=[
            RequirementItem(
                text="The bidder must qualify.",
                evidence=[EvidenceReference(paragraph_index=0, quote="must qualify")],
            )
        ]
    )


def override_provider(provider: object) -> None:
    main.app.dependency_overrides[main.get_llm_provider] = lambda: provider


def post_file(path: Path) -> object:
    with path.open("rb") as handle:
        return TestClient(main.app).post("/api/v1/tenders/analyze", files={"file": (path.name, handle)})


def test_docx_upload_parses_and_analyzes_with_injected_fake(sample_docx: Path) -> None:
    fake = FakeLLMProvider(analysis_response())
    override_provider(fake)
    response = post_file(sample_docx)
    assert response.status_code == 200
    result = TenderAnalysis.model_validate(response.json())
    assert result.qualification_requirements[0].text == "The bidder must qualify."
    assert fake.received_messages


def test_pdf_upload_analyzes_with_content_order_fallback(sample_pdf: Path) -> None:
    override_provider(FakeLLMProvider(analysis_response()))
    response = post_file(sample_pdf)
    assert response.status_code == 200
    assert "source_order_unavailable" in response.json()["warnings"]


def test_doc_dispatch_path_can_be_replaced(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    source = tmp_path / "legacy.doc"
    source.write_bytes(b"synthetic")
    seen: list[str] = []

    def fake_parse(path: Path, filename: str) -> ParsedDocument:
        seen.append(path.suffix)
        return ParsedDocument(filename=filename, file_type="doc")

    monkeypatch.setattr(main, "parse_document", fake_parse)
    override_provider(FakeLLMProvider(TenderChunkAnalysis()))
    response = post_file(source)
    assert response.status_code == 200
    assert seen == [".doc"]


def test_existing_parse_endpoint_still_works(sample_docx: Path) -> None:
    with sample_docx.open("rb") as handle:
        response = TestClient(main.app).post("/api/v1/documents/parse", files={"file": (sample_docx.name, handle)})
    assert response.status_code == 200
    assert response.json()["file_type"] == "docx"


def test_analyze_rejects_unsupported_and_empty_uploads() -> None:
    client = TestClient(main.app)
    unsupported = client.post("/api/v1/tenders/analyze", files={"file": ("note.txt", b"text")})
    empty = client.post("/api/v1/tenders/analyze", files={"file": ("empty.pdf", b"")})
    assert unsupported.status_code == 415
    assert empty.status_code == 422


def test_analyze_maps_size_limit(monkeypatch: pytest.MonkeyPatch, sample_docx: Path) -> None:
    async def too_large(*_args: object, **_kwargs: object) -> int:
        raise main.HTTPException(status_code=413, detail="synthetic size limit")

    monkeypatch.setattr(main, "stream_upload_to_path", too_large)
    override_provider(FakeLLMProvider(TenderChunkAnalysis()))
    assert post_file(sample_docx).status_code == 413


def test_analyze_maps_document_parse_failure(monkeypatch: pytest.MonkeyPatch, sample_docx: Path) -> None:
    monkeypatch.setattr(main, "parse_document", lambda *_args: (_ for _ in ()).throw(DocumentParseError("synthetic")))
    override_provider(FakeLLMProvider(TenderChunkAnalysis()))
    response = post_file(sample_docx)
    assert response.status_code == 422
    assert response.json()["detail"] == "Document could not be parsed."


@pytest.mark.parametrize(
    ("error", "status_code"),
    [
        (LLMConnectionError("synthetic"), 502),
        (LLMTimeoutError("synthetic"), 504),
        (LLMHTTPError(500), 502),
        (LLMStructuredOutputError("synthetic"), 502),
    ],
)
def test_provider_errors_are_mapped_without_secrets(sample_docx: Path, error: Exception, status_code: int) -> None:
    override_provider(FakeLLMProvider(error))
    response = post_file(sample_docx)
    assert response.status_code == status_code
    assert "synthetic" not in response.json()["detail"]
    assert "key" not in response.json()["detail"].casefold()


def test_all_chunks_failed_is_mapped(sample_docx: Path) -> None:
    override_provider(FakeLLMProvider(TenderAnalyzerError("synthetic")))
    assert post_file(sample_docx).status_code == 502


def test_provider_configuration_dependency_is_mapped_to_503(sample_docx: Path) -> None:
    def unavailable_provider() -> object:
        raise LLMConfigurationError("synthetic")

    main.app.dependency_overrides[main.get_llm_provider] = unavailable_provider
    response = post_file(sample_docx)
    assert response.status_code == 503


def test_no_content_pdf_returns_success_warning(tmp_path: Path) -> None:
    path = tmp_path / "drawing.pdf"
    pdf = pymupdf.open()
    pdf.new_page()
    pdf.save(path)
    pdf.close()
    fake = FakeLLMProvider(TenderChunkAnalysis())
    override_provider(fake)
    response = post_file(path)
    assert response.status_code == 200
    assert response.json()["warnings"] == ["document_has_no_extractable_content"]
    assert fake.received_messages == []


def test_temp_directory_is_cleaned_on_success_and_failure(monkeypatch: pytest.MonkeyPatch, sample_docx: Path, tmp_path: Path) -> None:
    paths: list[Path] = []

    class TrackingTemporaryDirectory:
        def __init__(self, prefix: str) -> None:
            self.path = tmp_path / f"{prefix}{len(paths)}"

        def __enter__(self) -> str:
            self.path.mkdir()
            paths.append(self.path)
            return str(self.path)

        def __exit__(self, *_args: object) -> None:
            shutil.rmtree(self.path)

    monkeypatch.setattr(main, "TemporaryDirectory", TrackingTemporaryDirectory)
    override_provider(FakeLLMProvider(TenderChunkAnalysis()))
    assert post_file(sample_docx).status_code == 200
    monkeypatch.setattr(main, "parse_document", lambda *_args: (_ for _ in ()).throw(DocumentParseError("synthetic")))
    assert post_file(sample_docx).status_code == 422
    assert paths and all(not path.exists() for path in paths)
