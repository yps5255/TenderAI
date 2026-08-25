from __future__ import annotations

import io
import shutil
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from pydantic import BaseModel

from backend.app import main
from backend.app.bid_analyzer.models import (
    BidCapabilityChunkAnalysis,
    BidCommercialServiceChunkAnalysis,
    BidCoreDocumentsChunkAnalysis,
    BidTechnicalChunkAnalysis,
    BidChunkQualificationMaterial,
    BidChunkTechnicalResponse,
    BidChunkTextItem,
    BidChunkPersonnelItem,
    BidChunkEquipmentItem,
    BidChunkExperienceItem,
    BidChunkCertificationItem,
    BidChunkSubmittedDocument,
    BidSourceEvidence,
)
from backend.app.bid_analyzer import BidAnalyzerError
from backend.app.llm.exceptions import (
    LLMConfigurationError,
    LLMConnectionError,
    LLMHTTPError,
    LLMStructuredOutputError,
    LLMTimeoutError,
)


@pytest.fixture(autouse=True)
def clear_dependency_overrides() -> None:
    main.app.dependency_overrides.clear()
    yield
    main.app.dependency_overrides.clear()


def synthetic_docx_bytes(*paragraphs: str) -> bytes:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def synthetic_bid_bytes() -> bytes:
    return synthetic_docx_bytes(
        "项目名称：测试模块化设备采购项目。项目编号：TEST-BID-2026-001。投标人：测试装备制造有限公司。",
        "资格材料：我公司已提供营业执照扫描件。我公司已提供特种设备生产许可证。",
        "招标要求设备额定功率不低于100kW。我方提供的设备额定功率为120kW。我方对此条技术要求无偏离。",
        "设备：数控切割设备2台。压力试验设备2台。人员：持证焊工12人。",
        "类似业绩：测试油气装备项目。客户为测试能源有限公司。合同金额500万元。完成时间2025年。",
        "投标报价：人民币280万元。交货承诺：合同签订后45日内完成供货。投标有效期：自投标截止之日起90日。",
        "认证：ISO9001质量管理体系认证。售后服务：设备验收后提供24个月质量保证。",
    )


def source(ref: str, quote: str | None = None) -> BidSourceEvidence:
    return BidSourceEvidence(source_ref=ref, quote=quote)


def grouped_responses() -> dict[type[BaseModel], BaseModel]:
    return {
        BidCoreDocumentsChunkAnalysis: BidCoreDocumentsChunkAnalysis(
            project_name="测试模块化设备采购项目",
            project_number="TEST-BID-2026-001",
            bidder="测试装备制造有限公司",
            qualification_materials=[
                BidChunkQualificationMaterial(name="营业执照", evidence=[source("P:1")]),
                BidChunkQualificationMaterial(name="特种设备生产许可证", evidence=[source("P:1")]),
            ],
            submitted_documents=[BidChunkSubmittedDocument(name="营业执照", evidence=[source("P:1")])],
        ),
        BidTechnicalChunkAnalysis: BidTechnicalChunkAnalysis(
            technical_responses=[BidChunkTechnicalResponse(response="120kW", evidence=[source("P:2")])],
            technical_solution=[BidChunkTextItem(text="模块化控制系统", evidence=[source("P:2")])],
        ),
        BidCapabilityChunkAnalysis: BidCapabilityChunkAnalysis(
            personnel=[BidChunkPersonnelItem(role="持证焊工", description="12人", evidence=[source("P:3")])],
            equipment=[BidChunkEquipmentItem(name="数控切割设备", quantity="2台", evidence=[source("P:3")])],
            experience_items=[BidChunkExperienceItem(project_name="测试油气装备项目", evidence=[source("P:4")])],
            certifications=[BidChunkCertificationItem(name="ISO9001", evidence=[source("P:4")])],
        ),
        BidCommercialServiceChunkAnalysis: BidCommercialServiceChunkAnalysis(
            bid_price="人民币280万元",
            delivery_commitment="合同签订后45日内完成供货",
            validity_period="90日",
            service_commitments=[BidChunkTextItem(text="24个月质量保证", evidence=[source("P:6")])],
        ),
    }


class RecordingProvider:
    def __init__(self, responses: dict[type[BaseModel], BaseModel] | None = None, failures: list[Exception] | None = None) -> None:
        self.responses = responses or grouped_responses()
        self.failures = list(failures or [])
        self.calls: list[type[BaseModel]] = []

    def generate(self, messages, response_model):
        self.calls.append(response_model)
        if self.failures:
            failure = self.failures.pop(0)
            if isinstance(failure, Exception):
                raise failure
        return response_model.model_validate(self.responses[response_model].model_dump())


def override_provider(provider: object) -> None:
    main.app.dependency_overrides[main.get_llm_provider] = lambda: provider


def post_docx(content: bytes, filename: str = "synthetic-bid.docx"):
    return TestClient(main.app).post("/api/v1/bids/analyze", files={"file": (filename, content)})


def test_bid_route_and_openapi_schema() -> None:
    schema = TestClient(main.app).get("/openapi.json").json()
    operation = schema["paths"]["/api/v1/bids/analyze"]["post"]
    assert operation["responses"]["200"]["content"]["application/json"]["schema"]["$ref"].endswith("/BidAnalysis")


def test_synthetic_docx_end_to_end_uses_parser_and_bid_analyzer() -> None:
    provider = RecordingProvider()
    override_provider(provider)
    response = post_docx(synthetic_bid_bytes())
    assert response.status_code == 200
    body = response.json()
    assert body["project_name"] == "测试模块化设备采购项目"
    assert body["project_number"] == "TEST-BID-2026-001"
    assert body["bidder"] == "测试装备制造有限公司"
    assert body["bid_price"] == "人民币280万元"
    assert body["technical_responses"][0]["evidence"][0].keys() >= {"paragraph_index", "page_number", "quote"}
    assert provider.calls


def test_empty_docx_returns_warning_without_provider_call() -> None:
    provider = RecordingProvider()
    override_provider(provider)
    response = post_docx(synthetic_docx_bytes())
    assert response.status_code == 200
    assert response.json()["warnings"] == ["document_has_no_extractable_content"]
    assert provider.calls == []


def test_unsupported_extension_returns_415() -> None:
    assert post_docx(b"text", "note.txt").status_code == 415


def test_parse_failure_returns_422(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(main, "parse_document", lambda *_args: (_ for _ in ()).throw(main.DocumentParseError("bad")))
    override_provider(RecordingProvider())
    response = post_docx(synthetic_bid_bytes())
    assert response.status_code == 422
    assert response.json()["detail"] == "Document could not be parsed."


def test_oversized_upload_returns_413(monkeypatch: pytest.MonkeyPatch) -> None:
    async def too_large(*_args, **_kwargs):
        raise main.HTTPException(status_code=413, detail="synthetic size limit")
    monkeypatch.setattr(main, "stream_upload_to_path", too_large)
    override_provider(RecordingProvider())
    assert post_docx(synthetic_bid_bytes()).status_code == 413


@pytest.mark.parametrize(
    ("error", "status_code"),
    [(LLMConnectionError("x"), 502), (LLMHTTPError(500), 502), (LLMTimeoutError("x"), 504),
     (LLMStructuredOutputError("x"), 502)],
)
def test_provider_errors_are_safe(error: Exception, status_code: int) -> None:
    override_provider(RecordingProvider(failures=[error for _ in range(8)]))
    response = post_docx(synthetic_docx_bytes("普通正文"))
    assert response.status_code == status_code
    assert "x" not in response.json()["detail"]


def test_provider_configuration_returns_503() -> None:
    main.app.dependency_overrides[main.get_llm_provider] = lambda: (_ for _ in ()).throw(LLMConfigurationError("secret"))
    response = post_docx(synthetic_docx_bytes("普通正文"))
    assert response.status_code == 503
    assert "secret" not in response.text


def test_recoverable_group_failure_returns_200_warning() -> None:
    provider = RecordingProvider(failures=[LLMStructuredOutputError("one"), LLMStructuredOutputError("two")])
    override_provider(provider)
    response = post_docx(synthetic_docx_bytes("普通正文"))
    assert response.status_code == 200
    assert any("group_analysis_failed" in warning for warning in response.json()["warnings"])


def test_structured_retry_recovery_returns_200_without_retry_warning() -> None:
    provider = RecordingProvider()
    original_generate = provider.generate
    attempts = {"count": 0}

    def generate(messages, response_model):
        if response_model is BidTechnicalChunkAnalysis and attempts["count"] == 0:
            attempts["count"] += 1
            provider.calls.append(response_model)
            raise LLMStructuredOutputError("synthetic")
        return original_generate(messages, response_model)

    provider.generate = generate
    override_provider(provider)
    response = post_docx(synthetic_docx_bytes("普通正文"))
    assert response.status_code == 200
    assert "retry" not in response.text.casefold()


def test_bid_analyzer_error_returns_502(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(main, "_analyze_parsed_bid_document", lambda *_args: (_ for _ in ()).throw(BidAnalyzerError("secret")))
    override_provider(RecordingProvider())
    response = post_docx(synthetic_docx_bytes("普通正文"))
    assert response.status_code == 502
    assert "secret" not in response.text


def test_temp_directory_cleanup_on_bid_parse_failure(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    paths: list[Path] = []

    class TrackingTemporaryDirectory:
        def __init__(self, prefix: str) -> None:
            self.path = tmp_path / f"{prefix}{len(paths)}"
        def __enter__(self) -> str:
            self.path.mkdir()
            paths.append(self.path)
            return str(self.path)
        def __exit__(self, *_args) -> None:
            shutil.rmtree(self.path)

    monkeypatch.setattr(main, "TemporaryDirectory", TrackingTemporaryDirectory)
    monkeypatch.setattr(main, "parse_document", lambda *_args: (_ for _ in ()).throw(main.DocumentParseError("bad")))
    override_provider(RecordingProvider())
    assert post_docx(synthetic_bid_bytes()).status_code == 422
    assert paths and all(not path.exists() for path in paths)
