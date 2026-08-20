from __future__ import annotations

from pathlib import Path

import pymupdf
import pytest
from docx import Document
from fastapi.testclient import TestClient

from backend.app.main import app
from backend.app.parser.document_parser import parse_document


@pytest.fixture()
def sample_docx(tmp_path: Path) -> Path:
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("第一段测试文本")
    document.add_paragraph("第二段测试文本")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(1, 0).text = "A2"
    table.cell(1, 1).text = "B2"
    document.save(path)
    return path


@pytest.fixture()
def sample_pdf(tmp_path: Path) -> Path:
    path = tmp_path / "sample.pdf"
    document = pymupdf.open()
    for text in ("Page one test text", "Page two test text"):
        page = document.new_page()
        page.insert_text((72, 72), text)
    document.save(path)
    document.close()
    return path


def test_health() -> None:
    response = TestClient(app).get("/health")
    assert response.status_code == 200
    assert response.json() == {"status": "ok"}


def test_parse_docx(sample_docx: Path) -> None:
    result = parse_document(sample_docx)
    assert result.file_type == "docx"
    assert [paragraph.text for paragraph in result.paragraphs] == ["第一段测试文本", "第二段测试文本"]
    assert result.tables[0].rows == [["A1", "B1"], ["A2", "B2"]]
    assert result.paragraphs[0].page_number is None


def test_parse_pdf(sample_pdf: Path) -> None:
    result = parse_document(sample_pdf)
    assert result.file_type == "pdf"
    assert result.page_count == 2
    assert len(result.pages) == 2
    assert "Page one test text" in result.pages[0].text
    assert result.pages[0].page_number == 1


def test_upload_docx(sample_docx: Path) -> None:
    with sample_docx.open("rb") as handle:
        response = TestClient(app).post(
            "/api/v1/documents/parse",
            files={"file": (sample_docx.name, handle, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
    assert response.status_code == 200
    assert response.json()["file_type"] == "docx"
    assert len(response.json()["tables"]) == 1


def test_upload_pdf(sample_pdf: Path) -> None:
    with sample_pdf.open("rb") as handle:
        response = TestClient(app).post(
            "/api/v1/documents/parse",
            files={"file": (sample_pdf.name, handle, "application/pdf")},
        )
    assert response.status_code == 200
    assert response.json()["page_count"] == 2


def test_reject_unsupported_file() -> None:
    response = TestClient(app).post(
        "/api/v1/documents/parse",
        files={"file": ("note.txt", b"text", "text/plain")},
    )
    assert response.status_code == 415


@pytest.mark.parametrize("filename, content", [("empty.pdf", b""), ("broken.docx", b"not a document")])
def test_bad_upload_does_not_crash(filename: str, content: bytes) -> None:
    response = TestClient(app).post(
        "/api/v1/documents/parse",
        files={"file": (filename, content)},
    )
    assert response.status_code == 422
